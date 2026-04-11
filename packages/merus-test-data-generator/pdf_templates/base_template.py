"""
Base PDF template — letterheads, headers, signatures, footers, lorem generators.
All specialized templates inherit from this.

@Developed & Documented by Glass Box Solutions, Inc. using human ingenuity and modern technology
"""

from __future__ import annotations

import random
from datetime import date
from io import BytesIO
from pathlib import Path
from typing import Any

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    HRFlowable,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


class BaseTemplate:
    """Base class for all PDF templates."""

    def __init__(self, case: Any):
        self.case = case
        self.styles = getSampleStyleSheet()
        self._register_custom_styles()

    def _register_custom_styles(self) -> None:
        self.styles.add(ParagraphStyle(
            name="Letterhead",
            fontSize=14,
            fontName="Helvetica-Bold",
            alignment=TA_CENTER,
            spaceAfter=2,
        ))
        self.styles.add(ParagraphStyle(
            name="LetterheadSub",
            fontSize=9,
            fontName="Helvetica",
            alignment=TA_CENTER,
            spaceAfter=1,
            textColor=colors.HexColor("#555555"),
        ))
        self.styles.add(ParagraphStyle(
            name="SectionHeader",
            fontSize=11,
            fontName="Helvetica-Bold",
            spaceAfter=6,
            spaceBefore=12,
        ))
        self.styles.add(ParagraphStyle(
            name="BodyText14",
            fontSize=10,
            fontName="Times-Roman",
            leading=14,
            alignment=TA_JUSTIFY,
            spaceAfter=6,
        ))
        self.styles.add(ParagraphStyle(
            name="DoubleSpaced",
            fontSize=12,
            fontName="Times-Roman",
            leading=24,
            alignment=TA_JUSTIFY,
            spaceAfter=6,
        ))
        self.styles.add(ParagraphStyle(
            name="MonoFindings",
            fontSize=9,
            fontName="Courier",
            leading=12,
            spaceAfter=4,
        ))
        self.styles.add(ParagraphStyle(
            name="SmallItalic",
            fontSize=8,
            fontName="Helvetica-Oblique",
            textColor=colors.HexColor("#888888"),
        ))
        self.styles.add(ParagraphStyle(
            name="RightAligned",
            fontSize=10,
            fontName="Times-Roman",
            alignment=TA_RIGHT,
        ))
        self.styles.add(ParagraphStyle(
            name="CenterBold",
            fontSize=11,
            fontName="Helvetica-Bold",
            alignment=TA_CENTER,
            spaceAfter=6,
        ))
        self.styles.add(ParagraphStyle(
            name="TableCell",
            fontSize=9,
            fontName="Helvetica",
            leading=11,
        ))
        self.styles.add(ParagraphStyle(
            name="Transcript",
            fontSize=10,
            fontName="Courier",
            leading=14,
            leftIndent=36,
        ))

    def generate(self, output_path: Path, doc_spec: Any) -> Path:
        """Dispatch to the correct format generator based on doc_spec.output_format."""
        fmt = getattr(doc_spec, "output_format", None)
        fmt_val = fmt.value if fmt is not None else "pdf"

        output_path.parent.mkdir(parents=True, exist_ok=True)

        if fmt_val == "eml":
            return self._generate_eml(output_path, doc_spec)
        elif fmt_val == "docx":
            return self._generate_docx(output_path, doc_spec)
        elif fmt_val == "scanned_pdf":
            return self._generate_scanned_pdf(output_path, doc_spec)
        else:
            return self._generate_pdf(output_path, doc_spec)

    def _generate_pdf(self, output_path: Path, doc_spec: Any) -> Path:
        """Render a native vector PDF using reportlab."""
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            leftMargin=0.75 * inch,
            rightMargin=0.75 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
        )
        story = self.build_story(doc_spec)
        doc.build(story, onFirstPage=self._add_footer, onLaterPages=self._add_footer)
        output_path.write_bytes(buffer.getvalue())
        return output_path

    def _generate_eml(self, output_path: Path, doc_spec: Any) -> Path:
        """Render an RFC 2822 .eml file from the document story."""
        import email.message
        from data.email_metadata import generate_email_headers

        story = self.build_story(doc_spec)
        body = self._story_to_plaintext(story)
        headers = generate_email_headers(
            subtype=doc_spec.subtype.value,
            case=self.case,
            doc_date=doc_spec.doc_date,
            subject=doc_spec.title,
        )

        msg = email.message.Message()
        for key, val in headers.items():
            msg[key] = val
        msg.set_payload(body, charset="utf-8")

        output_path.write_text(str(msg), encoding="utf-8")
        return output_path

    def _generate_docx(self, output_path: Path, doc_spec: Any) -> Path:
        """Render a Word .docx document from the document story."""
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from data.docx_styles import (
            STYLE_MAP,
            FIRM_LETTERHEAD,
            MONOSPACE_STYLES,
            DOUBLE_SPACED_STYLES,
            CENTER_ALIGNED_STYLES,
            RIGHT_ALIGNED_STYLES,
        )

        word_doc = Document()

        # Set Times New Roman 12pt as the default body font
        style = word_doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)

        # Firm letterhead in the header section
        section = word_doc.sections[0]
        header = section.header
        hdr_para = header.paragraphs[0]
        hdr_para.text = FIRM_LETTERHEAD["name"]
        hdr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if hdr_para.runs:
            hdr_para.runs[0].bold = True
            hdr_para.runs[0].font.size = Pt(14)
        header.add_paragraph(FIRM_LETTERHEAD["address"]).alignment = WD_ALIGN_PARAGRAPH.CENTER
        header.add_paragraph(FIRM_LETTERHEAD["phone"]).alignment = WD_ALIGN_PARAGRAPH.CENTER

        # CONFIDENTIAL footer with page number field
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "CONFIDENTIAL — Workers' Compensation Medical/Legal Record"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if footer_para.runs:
            footer_para.runs[0].font.size = Pt(8)

        # Convert story flowables to Word elements
        story = self.build_story(doc_spec)
        self._story_to_docx(story, word_doc, STYLE_MAP, MONOSPACE_STYLES,
                            DOUBLE_SPACED_STYLES, CENTER_ALIGNED_STYLES, RIGHT_ALIGNED_STYLES)

        word_doc.save(str(output_path))
        return output_path

    def _generate_scanned_pdf(self, output_path: Path, doc_spec: Any) -> Path:
        """Generate a native PDF then apply scan simulation artifacts."""
        import random as _random
        from pdf_templates.scan_simulator import simulate_scan

        # Generate clean native PDF first
        native_buf = BytesIO()
        native_doc = SimpleDocTemplate(
            native_buf,
            pagesize=letter,
            leftMargin=0.75 * inch,
            rightMargin=0.75 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
        )
        story = self.build_story(doc_spec)
        native_doc.build(story, onFirstPage=self._add_footer, onLaterPages=self._add_footer)

        # Apply scan simulation with a doc-specific seed derived from title hash
        seed = hash(doc_spec.title + str(doc_spec.doc_date)) & 0xFFFFFFFF
        rng = _random.Random(seed)
        scanned_bytes = simulate_scan(native_buf.getvalue(), rng, doc_date=doc_spec.doc_date)

        output_path.write_bytes(scanned_bytes)
        return output_path

    def build_story(self, doc_spec: Any) -> list:
        raise NotImplementedError

    # --- Interdocument coherence helpers (Phase 6) ---

    def _get_accumulator(self, doc_spec: Any) -> Any | None:
        """Return the CaseContextAccumulator for this case, or None if not wired.

        Templates should always call this rather than accessing doc_spec.context
        directly. Returning None means the pipeline hasn't injected the accumulator
        (e.g. unit tests or legacy pipeline code) — templates must degrade gracefully.
        """
        return doc_spec.context.get("_accumulator") if doc_spec.context else None

    def _format_cross_references(self, doc_spec: Any, max_refs: int = 3) -> str:
        """Return a formatted cross-reference sentence from the accumulator.

        Returns an empty string if no accumulator is available or no prior
        documents have been generated yet.
        """
        acc = self._get_accumulator(doc_spec)
        if acc is None:
            return ""
        return acc.get_cross_reference(max_refs=max_refs)

    # --- Format conversion helpers ---

    def _story_to_plaintext(self, story: list) -> str:
        """Extract readable plain text from a reportlab story.

        Walks Paragraph, Table, Spacer, and HRFlowable flowables and converts
        them to plain text suitable for an email body.
        """
        import re
        from reportlab.platypus import Paragraph, Table, Spacer, HRFlowable

        lines: list[str] = []
        for flowable in story:
            if isinstance(flowable, Paragraph):
                # Strip all reportlab/HTML tags
                text = re.sub(r"<[^>]+>", "", flowable.text or "")
                text = text.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
                lines.append(text.strip())
            elif isinstance(flowable, Table):
                # Format table rows as tab-separated text
                for row in (flowable._cellvalues if hasattr(flowable, "_cellvalues") else []):
                    cells = []
                    for cell in row:
                        if isinstance(cell, Paragraph):
                            cell_text = re.sub(r"<[^>]+>", "", cell.text or "")
                        else:
                            cell_text = str(cell) if cell else ""
                        cells.append(cell_text.strip())
                    lines.append("\t".join(cells))
            elif isinstance(flowable, Spacer):
                lines.append("")
            elif isinstance(flowable, HRFlowable):
                lines.append("-" * 60)

        return "\n".join(lines)

    def _story_to_docx(
        self,
        story: list,
        word_doc: Any,
        style_map: dict,
        monospace_styles: frozenset,
        double_spaced_styles: frozenset,
        center_styles: frozenset,
        right_styles: frozenset,
    ) -> None:
        """Convert a reportlab story into python-docx paragraphs and tables.

        Mutates word_doc in place by appending elements.
        """
        import re
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from reportlab.platypus import Paragraph, Table, Spacer, HRFlowable

        def strip_tags(text: str) -> str:
            return re.sub(r"<[^>]+>", "", text or "").strip()

        def is_bold(text: str) -> bool:
            return bool(re.search(r"<b>", text or "", re.IGNORECASE))

        def is_italic(text: str) -> bool:
            return bool(re.search(r"<i>|<em>", text or "", re.IGNORECASE))

        for flowable in story:
            if isinstance(flowable, Paragraph):
                style_name = flowable.style.name if hasattr(flowable, "style") else "Normal"
                docx_style_name, force_bold, force_italic, font_pt = style_map.get(
                    style_name, ("Normal", False, False, None)
                )

                para = word_doc.add_paragraph(style=docx_style_name)
                raw_text = flowable.text or ""
                run = para.add_run(strip_tags(raw_text))

                if force_bold or is_bold(raw_text):
                    run.bold = True
                if force_italic or is_italic(raw_text):
                    run.italic = True
                if font_pt:
                    run.font.size = Pt(font_pt)
                if style_name in monospace_styles:
                    run.font.name = "Courier New"
                if style_name in center_styles:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif style_name in right_styles:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if style_name in double_spaced_styles:
                    para.paragraph_format.line_spacing = Pt(24)

            elif isinstance(flowable, Table):
                cell_values = getattr(flowable, "_cellvalues", [])
                if not cell_values:
                    continue
                rows = len(cell_values)
                cols = max(len(r) for r in cell_values) if cell_values else 1
                tbl = word_doc.add_table(rows=rows, cols=cols)
                tbl.style = "Table Grid"
                for r_idx, row in enumerate(cell_values):
                    for c_idx, cell in enumerate(row):
                        if c_idx >= cols:
                            break
                        if isinstance(cell, Paragraph):
                            cell_text = strip_tags(cell.text)
                        else:
                            cell_text = str(cell) if cell else ""
                        tbl.rows[r_idx].cells[c_idx].text = cell_text

            elif isinstance(flowable, Spacer):
                word_doc.add_paragraph("")

            elif isinstance(flowable, HRFlowable):
                # Add a paragraph with a bottom border to simulate a horizontal rule
                para = word_doc.add_paragraph()
                pPr = para._p.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "6")
                bottom.set(qn("w:space"), "1")
                bottom.set(qn("w:color"), "333333")
                pBdr.append(bottom)
                pPr.append(pBdr)

    # --- Reusable components ---

    def make_letterhead(self, org: str, address: str, phone: str) -> list:
        return [
            Paragraph(org, self.styles["Letterhead"]),
            Paragraph(address, self.styles["LetterheadSub"]),
            Paragraph(phone, self.styles["LetterheadSub"]),
            HRFlowable(width="100%", thickness=1, color=colors.HexColor("#333333")),
            Spacer(1, 12),
        ]

    def make_patient_header(self) -> list:
        a = self.case.applicant
        data = [
            ["Patient:", a.full_name, "DOB:", a.date_of_birth.strftime("%m/%d/%Y")],
            ["SSN (last 4):", f"XXX-XX-{a.ssn_last_four}", "Phone:", a.phone],
        ]
        t = Table(data, colWidths=[80, 200, 60, 140])
        t.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
            ("FONTNAME", (2, 0), (2, -1), "Helvetica-Bold"),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ("TOPPADDING", (0, 0), (-1, -1), 2),
        ]))
        return [t, Spacer(1, 6)]

    def make_claim_reference_block(self) -> list:
        inj = self.case.injuries[0]
        ins = self.case.insurance
        emp = self.case.employer
        data = [
            ["ADJ Number:", inj.adj_number, "Claim #:", ins.claim_number],
            ["Date of Injury:", inj.date_of_injury.strftime("%m/%d/%Y"), "Employer:", emp.company_name],
            ["Injury Type:", inj.injury_type.value.replace("_", " ").title(), "Carrier:", ins.carrier_name],
            ["Body Parts:", ", ".join(inj.body_parts), "", ""],
        ]
        t = Table(data, colWidths=[85, 175, 65, 155])
        t.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
            ("FONTNAME", (2, 0), (2, -1), "Helvetica-Bold"),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ("TOPPADDING", (0, 0), (-1, -1), 2),
            ("BOX", (0, 0), (-1, -1), 0.5, colors.grey),
            ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f0f0f0")),
            ("BACKGROUND", (2, 0), (2, -1), colors.HexColor("#f0f0f0")),
        ]))
        return [t, Spacer(1, 12)]

    def make_signature_block(self, name: str, title: str, license_num: str = "") -> list:
        elements = [
            Spacer(1, 30),
            HRFlowable(width="40%", thickness=0.5, color=colors.black),
            Paragraph(name, self.styles["BodyText14"]),
            Paragraph(title, self.styles["SmallItalic"]),
        ]
        if license_num:
            elements.append(Paragraph(f"License: {license_num}", self.styles["SmallItalic"]))
        return elements

    def make_date_line(self, label: str, d: date) -> Paragraph:
        return Paragraph(f"<b>{label}:</b> {d.strftime('%B %d, %Y')}", self.styles["BodyText14"])

    def make_section(self, title: str, content: str) -> list:
        return [
            Paragraph(title, self.styles["SectionHeader"]),
            Paragraph(content, self.styles["BodyText14"]),
        ]

    def make_hr(self) -> HRFlowable:
        return HRFlowable(width="100%", thickness=0.5, color=colors.grey)

    def _add_footer(self, canvas, doc) -> None:
        # --- Watermark (top of page, OCR-resistant) ---
        canvas.saveState()
        canvas.setFillAlpha(0.07)  # 7% opacity — below OCR binarization threshold
        canvas.setFillColor(colors.HexColor("#888888"))
        canvas.setFont("Helvetica-Bold", 16)
        text = "GBS Generated"
        text_width = canvas.stringWidth(text, "Helvetica-Bold", 16)
        x = (letter[0] - text_width) / 2  # Centered horizontally
        y = letter[1] - 0.45 * inch       # Top margin area
        canvas.drawString(x, y, text)
        canvas.restoreState()

        # --- Footer ---
        canvas.saveState()
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(colors.HexColor("#999999"))
        canvas.drawString(
            0.75 * inch, 0.4 * inch,
            "CONFIDENTIAL — Workers' Compensation Medical/Legal Record"
        )
        canvas.drawRightString(
            letter[0] - 0.75 * inch, 0.4 * inch,
            f"Page {doc.page}"
        )
        canvas.restoreState()

    # --- Lorem generators ---

    @staticmethod
    def lorem_medical(sentences: int = 3) -> str:
        pool = [
            "Patient presents with complaints of persistent pain and limited range of motion.",
            "Physical examination reveals tenderness to palpation over the affected area with guarding.",
            "Neurological examination demonstrates intact sensation and motor function bilaterally.",
            "Range of motion is restricted in flexion and extension compared to contralateral side.",
            "No signs of acute distress; patient is alert and oriented x3.",
            "Muscle strength testing reveals 4/5 weakness in the affected extremity.",
            "Deep tendon reflexes are 2+ and symmetric bilaterally.",
            "Straight leg raise test is positive on the affected side at 45 degrees.",
            "Spurling's test reproduces radicular symptoms with lateral flexion.",
            "Gait is antalgic with reduced stride length on the affected side.",
            "Imaging studies correlate with clinical findings of degenerative changes.",
            "Patient reports pain level of 6/10 at rest, increasing to 8/10 with activity.",
            "Swelling and ecchymosis are noted over the affected joint.",
            "Crepitus is palpable with active range of motion testing.",
            "Provocative testing reproduces the patient's primary complaint.",
            "Sensory examination reveals diminished light touch in the affected dermatome.",
            "Patient demonstrates functional limitations consistent with objective findings.",
            "No evidence of malingering or symptom magnification on Waddell's testing.",
            # Expanded pool — general clinical sentences
            "Motor strength is graded 4+/5 in the affected muscle group, a slight decrease from prior examination.",
            "Inspection reveals no erythema, ecchymosis, or open wounds in the affected area.",
            "Joint stability testing reveals no instability; ligamentous structures appear intact.",
            "The patient demonstrates appropriate effort during examination; no inconsistencies noted.",
            "Palpation of the paravertebral musculature reveals bilateral muscle spasm, left greater than right.",
            "Peripheral pulses are intact and symmetric in the upper and lower extremities.",
            "Skin examination reveals no trophic changes, suggesting adequate vascular supply.",
            "The patient reports difficulty with prolonged sitting beyond 30 minutes due to increased symptoms.",
            "Coordination testing is within normal limits; no dysmetria or tremor observed.",
            "Active range of motion produces pain at end range; passive motion exceeds active by approximately 10 degrees.",
            "The patient's subjective complaints are consistent with the objective examination findings.",
            "No atrophy is noted on visual inspection or circumferential measurements of the extremities.",
            "Reflexes are brisk at 3+ at the patella bilaterally but without clonus.",
            "The patient demonstrates a slow, deliberate gait pattern with slightly shortened stride length.",
            "Cervical spine examination reveals restricted rotation to the right with reproduction of concordant symptoms.",
            "Upper extremity motor examination is 5/5 throughout; grip strength is symmetric bilaterally.",
            "Lower extremity sensation is intact to light touch, pinprick, and vibration in all dermatomes tested.",
            "The patient's body habitus is noted; BMI is in the overweight range which may be a contributing factor.",
            "No lymphadenopathy or thyromegaly is appreciated on cervical examination.",
            "Cardiovascular examination reveals regular rate and rhythm without murmurs, rubs, or gallops.",
            "Respiratory examination reveals clear lung fields bilaterally with no adventitious sounds.",
            "Abdominal examination is benign; no tenderness, guarding, or masses appreciated.",
            "The patient was examined in both seated and supine positions to assess for positional variance.",
            "Provocative maneuvers are consistently positive, supporting the clinical diagnosis.",
            "The examiner notes that findings are reproducible across repeated testing during this examination.",
            "Functional capacity is estimated at the sedentary to light physical demand level based on clinical assessment.",
            "Patient demonstrates difficulty transitioning from seated to standing position, requiring use of armrests.",
            "Bilateral hand grip strength is reduced compared to age-matched normative values.",
            "Examination of the contralateral extremity reveals normal findings, serving as an internal control.",
            "No signs of complex regional pain syndrome (CRPS) are observed; no allodynia, color changes, or temperature asymmetry.",
            "The patient's effort during manual muscle testing is assessed as appropriate and consistent.",
            "Tandem gait is performed without difficulty; Romberg test is negative.",
            "The patient is able to remove shoes and socks independently but reports difficulty with sock application.",
            "Observation of spontaneous movements during the interview reveals occasional guarding with position changes.",
            "Vital signs are within normal limits: blood pressure, heart rate, and respiratory rate are unremarkable.",
            "The patient's pain behaviors during examination are proportionate to the reported complaint.",
            "No surgical scars are observed in the region of complaint.",
            "The patient maintains good eye contact and appears to provide a reliable history.",
            "Cognitive function appears grossly intact during the clinical interview.",
            "Manual muscle testing of the core stabilizers reveals mild weakness on the affected side.",
            "The neurovascular status of the affected extremity is intact distally.",
            "Skin turgor and capillary refill are normal in the affected extremity.",
        ]
        return " ".join(random.sample(pool, min(sentences, len(pool))))

    @staticmethod
    def lorem_legal(sentences: int = 3) -> str:
        pool = [
            "Applicant sustained industrial injury arising out of and in the course of employment.",
            "The injury is compensable under Labor Code sections 3600 et seq.",
            "Defendant carrier has accepted liability for the claimed body parts.",
            "Medical treatment has been provided pursuant to Labor Code section 4600.",
            "Temporary disability benefits have been paid at the statutory rate.",
            "The matter is now at issue regarding permanent disability and need for future medical treatment.",
            "Applicant's treating physician has indicated that the condition is permanent and stationary.",
            "A Qualified Medical Evaluator has been appointed to address disputed medical issues.",
            "Discovery is ongoing to determine the extent of permanent disability.",
            "The parties have been unable to resolve the matter informally.",
            "Applicant reserves the right to claim additional body parts as supported by medical evidence.",
            "The employer had knowledge of the injury and timely reported same to the carrier.",
            "All statutory notices have been provided to the applicant as required by law.",
            "The matter is set for hearing before the Workers' Compensation Appeals Board.",
            # Expanded pool — legal sentences
            "Pursuant to Labor Code section 4660.1, permanent disability shall be determined using the AMA Guides, Fifth Edition.",
            "The applicant's average weekly earnings at the time of injury establish the applicable indemnity rate.",
            "Defendant disputes the nature and extent of permanent disability as alleged by applicant.",
            "The treating physician's opinion on causation and apportionment is substantial medical evidence per LC §4061.",
            "Applicant contends that the industrial injury is a contributing cause of the need for medical treatment.",
            "The parties stipulate to the date of injury, employer-employee relationship, and jurisdiction of the WCAB.",
            "Defendant reserves the right to present evidence of apportionment under LC §4663 and §4664.",
            "The matter has been designated as a priority conference per 8 CCR §10205.10.",
            "Applicant has complied with all discovery requests and statutory reporting requirements.",
            "The defense medical evaluator's opinions are controverted by the treating physician's findings.",
            "Penalties and interest may be applicable if benefits are not timely provided per LC §5814.",
            "The applicant's claim for self-procured medical treatment is supported by documentation of the carrier's failure to provide timely care.",
            "Defendant has issued a Notice of Offer of Modified or Alternative Work per DWC Form AD 10118.",
            "Applicant seeks reimbursement of reasonable litigation costs and expenses per 8 CCR §10451.",
            "The matter is ripe for adjudication of all disputed issues.",
            "Pursuant to LC §4062.2, the QME panel was properly requested and the evaluator selected.",
            "Applicant's vocational expert has determined diminished future earning capacity attributable to the industrial injury.",
            "The utilization review process has been exhausted and the matter is appropriate for IMR per LC §4610.5.",
            "Defendant contends that the applicant has reached maximum medical improvement with no ratable permanent disability.",
            "The applicant's entitlement to supplemental job displacement benefit (SJDB) voucher is established per LC §4658.7.",
            "Applicant has been temporarily totally disabled and unable to return to any employment since the date of injury.",
            "The employer failed to offer regular, modified, or alternative work within 60 days of the P&S date per LC §4658.6.",
            "Applicant's claim for future medical treatment is supported by the QME's recommendation for ongoing care.",
            "The parties are engaged in good faith settlement negotiations pursuant to the mandatory settlement conference.",
            "Discovery cut-off is approaching; all outstanding discovery must be completed prior to the trial date.",
            "The applicant's testimony at trial will establish the mechanism of injury and resulting disability.",
            "Medical evidence establishes that the industrial injury is the predominant cause of the applicant's psychiatric condition per LC §3208.3.",
            "The defense has not timely responded to applicant's discovery requests, and sanctions may be sought.",
            "The applicant is entitled to reasonable attorney's fees pursuant to LC §4903.1.",
            "Settlement discussions have been ongoing; a mandatory settlement conference is scheduled.",
            "The case involves multiple body parts with potential combined impairment exceeding 50% whole person.",
            "Applicant's case-in-chief will include testimony, medical records, and expert medical opinions.",
        ]
        return " ".join(random.sample(pool, min(sentences, len(pool))))

    @staticmethod
    def lorem_correspondence(sentences: int = 3) -> str:
        pool = [
            "Please find enclosed the documents referenced in our prior correspondence.",
            "We are writing to advise you of the current status of this workers' compensation claim.",
            "Kindly review the attached and respond at your earliest convenience.",
            "This letter serves as formal notice regarding the above-referenced matter.",
            "We have completed our review of the medical records and provide the following summary.",
            "Please do not hesitate to contact our office should you require additional information.",
            "We look forward to resolving this matter in an expeditious manner.",
            "Enclosed please find copies of all relevant documentation for your records.",
            "Thank you for your prompt attention to this matter.",
            "We respectfully request a response within thirty (30) days of the date of this letter.",
            # Expanded pool — correspondence sentences
            "As discussed during our telephone conference, the following issues remain outstanding.",
            "We are in receipt of your letter dated the above date and respond as follows.",
            "Please be advised that our client's position on this matter remains unchanged.",
            "This office represents the applicant in the above-referenced workers' compensation matter.",
            "We have been retained by the carrier to defend the employer in this claim.",
            "Pursuant to your request, we are providing the following updated medical records.",
            "We wish to bring to your attention the recent QME report which supports our client's position.",
            "Please confirm receipt of the enclosed documents and advise of any additional information needed.",
            "We note that the statutory deadline for response is approaching and request your timely reply.",
            "The applicant has authorized the release of medical records for the purpose of this claim.",
            "We are scheduling a mandatory settlement conference and request available dates from all parties.",
            "This correspondence is without prejudice to any rights, claims, or defenses of the parties.",
            "Please direct all future correspondence regarding this matter to the undersigned.",
            "We have reviewed the utilization review determination and intend to file an IMR request.",
            "The enclosed medical report should be reviewed in conjunction with the previously submitted records.",
            "We request that you authorize the recommended treatment as outlined in the attached Request for Authorization.",
            "The applicant continues to experience symptoms and requires ongoing medical treatment.",
            "We acknowledge receipt of the Notice of Offer of Work and will respond within the statutory timeframe.",
            "Please provide updated benefit payment records for our review and reconciliation.",
            "We request a copy of the complete claims file for our records.",
            "The enclosed invoice represents costs incurred in connection with the applicant's medical treatment.",
            "We anticipate that this matter can be resolved through informal negotiation.",
            "Your cooperation in facilitating the applicant's return to work is appreciated.",
            "We are providing this status update as required under our reporting obligations.",
            "The applicant's current work status and restrictions are detailed in the attached physician report.",
        ]
        return " ".join(random.sample(pool, min(sentences, len(pool))))

    # --- Specialty-aware content helpers ---

    def specialty_exam_findings(self, count: int = 8) -> str:
        """Get specialty-specific exam findings based on physician specialty."""
        from data.content_pools import get_exam_findings
        specialty = (
            getattr(self.case, 'qme_physician', None) and self.case.qme_physician.specialty
        ) or self.case.treating_physician.specialty
        body_parts = self.case.injuries[0].body_parts if self.case.injuries else []
        return get_exam_findings(specialty, body_parts, count)

    def rom_findings_table(self) -> Table:
        """Generate a ROM measured-vs-normal table for case body parts."""
        from data.content_pools import get_rom_table
        body_parts = self.case.injuries[0].body_parts if self.case.injuries else []
        rom_data = get_rom_table(body_parts)
        if not rom_data:
            return Spacer(1, 0)  # No ROM data available

        table_data = [["Body Part", "Movement", "Normal", "Measured", "Deficit"]]
        for row in rom_data:
            table_data.append([
                row["body_part"].title(),
                row["movement"],
                row["normal"],
                row["measured"],
                row["deficit"],
            ])

        t = Table(table_data, colWidths=[1.3 * inch, 1.5 * inch, 0.8 * inch, 0.9 * inch, 0.7 * inch])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2c3e50")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 8),
            ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
            ("FONTSIZE", (0, 1), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            *[("BACKGROUND", (0, i), (-1, i), colors.HexColor("#f8f9fa"))
              for i in range(2, len(table_data), 2)],
        ]))
        return t

    def impairment_rating_section(self) -> list:
        """Generate full AMA Guides impairment narrative as reportlab elements."""
        from data.ama_guides_content import generate_impairment_narrative
        body_parts = self.case.injuries[0].body_parts if self.case.injuries else []
        specialty = (
            getattr(self.case, 'qme_physician', None) and self.case.qme_physician.specialty
        ) or self.case.treating_physician.specialty
        apportionment_pct = random.choice([0, 0, 0, 10, 15, 20, 25])
        narrative, total_wpi, ratings = generate_impairment_narrative(
            body_parts, specialty, apportionment_pct
        )
        elements = [
            Paragraph("<b>IMPAIRMENT RATING — AMA GUIDES 5TH EDITION</b>", self.styles["SectionHeader"]),
            Spacer(1, 0.1 * inch),
        ]
        for para in narrative.split("\n"):
            if para.strip():
                elements.append(Paragraph(para, self.styles["BodyText14"]))
        elements.append(Spacer(1, 0.15 * inch))
        return elements

    def medical_record_review_list(self, doc_date) -> str:
        """Generate a detailed list of reviewed medical records."""
        from data.content_pools import get_record_review_items
        body_parts = self.case.injuries[0].body_parts if self.case.injuries else []
        position = self.case.employer.position
        items = get_record_review_items(body_parts, position, count=random.randint(12, 20))
        return "\n".join([f"• {item}" for item in items])

    def mtus_rationale(self, decision_type: str, count: int = 4) -> str:
        """Generate MTUS-citation-rich clinical rationale."""
        from data.content_pools import get_clinical_rationale, get_mtus_citations
        body_parts = self.case.injuries[0].body_parts if self.case.injuries else []
        rationale = get_clinical_rationale(decision_type, body_parts, count)
        citations = get_mtus_citations(body_parts, count=3)
        citations_str = "\n".join([f"• {c}" for c in citations])
        return f"{rationale}\n\n<b>Applicable Guidelines:</b>\n{citations_str}"
