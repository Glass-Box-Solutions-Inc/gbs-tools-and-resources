"""
Tests for Phase 3: .docx Word document generation.

Verifies that attorney work-product subtypes produce valid Word documents
with firm letterhead, correct styles, and no PDF/binary content.

@Developed & Documented by Glass Box Solutions, Inc. using human ingenuity and modern technology
"""

from __future__ import annotations

import tempfile
from pathlib import Path

import pytest
from docx import Document as DocxDocument

from data.models import DocumentSubtype, OutputFormat


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _generate_docx(template_cls, case, doc_spec) -> Path:
    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp) / "test.docx"
        template = template_cls(case)
        template.generate(out, doc_spec)
        assert out.exists(), "DOCX file was not created"
        assert out.stat().st_size > 0, "DOCX file is empty"
        # Copy bytes out before tmp is cleaned up
        contents = out.read_bytes()
    # Write to a stable temp file for inspection
    import tempfile as tf
    f = tf.NamedTemporaryFile(suffix=".docx", delete=False)
    f.write(contents)
    f.close()
    return Path(f.name)


# ---------------------------------------------------------------------------
# Settlement Valuation Memo — canonical DOCX template test
# ---------------------------------------------------------------------------

class TestSettlementMemoDocx:
    @pytest.fixture()
    def memo_spec(self, make_document_spec):
        return make_document_spec(
            subtype=DocumentSubtype.SETTLEMENT_VALUATION_MEMO,
            output_format=OutputFormat.DOCX,
            title="Settlement Valuation Memo",
        )

    def test_produces_valid_docx_file(self, sample_case, memo_spec):
        from pdf_templates.summaries.settlement_memo import SettlementMemo
        path = _generate_docx(SettlementMemo, sample_case, memo_spec)
        doc = DocxDocument(str(path))
        assert doc is not None

    def test_file_has_paragraphs(self, sample_case, memo_spec):
        from pdf_templates.summaries.settlement_memo import SettlementMemo
        path = _generate_docx(SettlementMemo, sample_case, memo_spec)
        doc = DocxDocument(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert len(text.strip()) > 50, "DOCX has no meaningful text content"

    def test_not_a_pdf_file(self, sample_case, memo_spec):
        from pdf_templates.summaries.settlement_memo import SettlementMemo
        path = _generate_docx(SettlementMemo, sample_case, memo_spec)
        magic = path.read_bytes()[:4]
        assert magic != b"%PDF", "DOCX file starts with PDF magic bytes"

    def test_is_valid_zip_container(self, sample_case, memo_spec):
        """OOXML .docx files are ZIP archives."""
        import zipfile
        from pdf_templates.summaries.settlement_memo import SettlementMemo
        path = _generate_docx(SettlementMemo, sample_case, memo_spec)
        assert zipfile.is_zipfile(str(path)), "DOCX file is not a valid ZIP/OOXML"


class TestMedicalChronologyDocx:
    @pytest.fixture()
    def chrono_spec(self, make_document_spec):
        return make_document_spec(
            subtype=DocumentSubtype.MEDICAL_CHRONOLOGY_TIMELINE,
            output_format=OutputFormat.DOCX,
            title="Medical Chronology / Timeline",
        )

    def test_produces_valid_docx(self, sample_case, chrono_spec):
        from pdf_templates.summaries.medical_chronology import MedicalChronology
        path = _generate_docx(MedicalChronology, sample_case, chrono_spec)
        doc = DocxDocument(str(path))
        assert doc is not None

    def test_has_content(self, sample_case, chrono_spec):
        from pdf_templates.summaries.medical_chronology import MedicalChronology
        path = _generate_docx(MedicalChronology, sample_case, chrono_spec)
        doc = DocxDocument(str(path))
        text = " ".join(p.text for p in doc.paragraphs)
        assert len(text.strip()) > 20


# ---------------------------------------------------------------------------
# Header / footer verification
# ---------------------------------------------------------------------------

class TestDocxHeaderFooter:
    @pytest.fixture()
    def memo_spec(self, make_document_spec):
        return make_document_spec(
            subtype=DocumentSubtype.SETTLEMENT_VALUATION_MEMO,
            output_format=OutputFormat.DOCX,
        )

    def test_header_contains_firm_name(self, sample_case, memo_spec):
        from pdf_templates.summaries.settlement_memo import SettlementMemo
        from data.docx_styles import FIRM_LETTERHEAD
        path = _generate_docx(SettlementMemo, sample_case, memo_spec)
        doc = DocxDocument(str(path))
        header_text = " ".join(
            p.text for section in doc.sections for p in section.header.paragraphs
        )
        assert FIRM_LETTERHEAD["name"] in header_text, (
            f"Firm name not in header. Got: {header_text!r}"
        )

    def test_footer_contains_confidential(self, sample_case, memo_spec):
        from pdf_templates.summaries.settlement_memo import SettlementMemo
        path = _generate_docx(SettlementMemo, sample_case, memo_spec)
        doc = DocxDocument(str(path))
        footer_text = " ".join(
            p.text for section in doc.sections for p in section.footer.paragraphs
        )
        assert "CONFIDENTIAL" in footer_text, (
            f"CONFIDENTIAL not in footer. Got: {footer_text!r}"
        )
